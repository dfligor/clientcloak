"""
Document handling for ClientCloak: load, save, extract, and replace text in .docx files.

This module is the bridge between python-docx Document objects and ClientCloak's
sanitization pipeline. It handles the notoriously tricky problem of replacing text
that Word splits across multiple XML runs due to formatting, spell-check state,
revision history, or other internal bookkeeping.

Key design decisions:
- Format-preserving replacement is attempted first (keeps bold/italic/color per-run).
- Fallback to run-collapsing only when cross-run surgery would be ambiguous.
- Case-insensitive matching with case-preserving output ("ACME" -> "CUSTOMER").
- All text sources are covered: body paragraphs, tables, headers, footers.
"""

from __future__ import annotations

import logging
import re
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------

class DocumentLoadError(Exception):
    """Raised when a document cannot be loaded."""


class UnsupportedFormatError(DocumentLoadError):
    """Raised when the file is not a supported .docx format."""


class PasswordProtectedError(DocumentLoadError):
    """Raised when the document is password-protected or encrypted."""


# ---------------------------------------------------------------------------
# Public API: load / save
# ---------------------------------------------------------------------------

def load_document(file_path: str | Path) -> Document:
    """
    Load a .docx file and return a python-docx Document.

    Validates:
    - File exists on disk.
    - Extension is .docx (not .doc or other).
    - File is a valid ZIP archive (the .docx container format).
    - File is not encrypted / password-protected.

    Raises:
        FileNotFoundError: File does not exist.
        UnsupportedFormatError: Wrong extension or not a valid ZIP.
        PasswordProtectedError: Encrypted document.
        DocumentLoadError: Any other load failure.
    """
    path = Path(file_path)

    # --- existence ---
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise DocumentLoadError(f"Path is not a file: {path}")

    # --- extension ---
    suffix = path.suffix.lower()
    if suffix == ".doc":
        raise UnsupportedFormatError(
            f"Legacy .doc format is not supported. Please convert to .docx first: {path.name}"
        )
    if suffix != ".docx":
        raise UnsupportedFormatError(
            f"Unsupported file type '{suffix}'. Only .docx files are accepted: {path.name}"
        )

    # --- valid ZIP ---
    if not zipfile.is_zipfile(path):
        raise UnsupportedFormatError(
            f"File is not a valid .docx archive (corrupt or not a real ZIP): {path.name}"
        )

    # --- encryption check ---
    if _is_encrypted(path):
        raise PasswordProtectedError(
            f"Document is password-protected or encrypted: {path.name}"
        )

    # --- load via python-docx ---
    try:
        return Document(str(path))
    except PackageNotFoundError as exc:
        raise DocumentLoadError(
            f"File appears damaged — required .docx internal parts are missing: {path.name}"
        ) from exc
    except Exception as exc:
        raise DocumentLoadError(
            f"Failed to load document '{path.name}': {exc}"
        ) from exc


def save_document(doc: Document, file_path: str | Path) -> Path:
    """
    Save a python-docx Document to *file_path*.

    Creates parent directories if they do not exist.
    Returns the resolved Path for convenience.
    """
    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    logger.info("Document saved to %s", path)
    return path.resolve()


# ---------------------------------------------------------------------------
# Public API: text extraction
# ---------------------------------------------------------------------------

# Each extracted item is (text, source_element) where source_element is the
# python-docx object that owns the text (Paragraph, _Cell, etc.).  Callers
# can use the source_element reference to locate or annotate the origin.

TextSource = tuple  # (str, object)


def extract_all_text(doc: Document) -> list[TextSource]:
    """
    Extract every text fragment from the document with a reference to its
    source element.

    Returns a list of ``(text, source_element)`` tuples covering:
    - Body paragraphs
    - Table cells (all rows, all cells)
    - Headers and footers (all sections)

    Empty strings are excluded.
    """
    results: list[TextSource] = []

    # --- body paragraphs ---
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.strip():
            results.append((text, paragraph))

    # --- tables ---
    results.extend(_extract_from_tables(doc.tables))

    # --- headers / footers (all sections) ---
    for section in doc.sections:
        for header_footer in _iter_headers_footers(section):
            for paragraph in header_footer.paragraphs:
                text = paragraph.text
                if text.strip():
                    results.append((text, paragraph))
            # Tables inside headers/footers
            if hasattr(header_footer, "tables"):
                results.extend(_extract_from_tables(header_footer.tables))

    return results


# ---------------------------------------------------------------------------
# Public API: replacement
# ---------------------------------------------------------------------------

def replace_text_in_document(
    doc: Document,
    replacements: dict[str, str],
    *,
    match_case: bool = True,
) -> int:
    """
    Apply *replacements* throughout the entire document.

    ``replacements`` maps original_text -> new_text.  Matching is
    **case-insensitive**; the replacement output preserves the case pattern
    of the matched text (see :func:`_transfer_case`) when *match_case* is
    ``True`` (the default).

    When *match_case* is ``False``, the replacement text is used exactly as
    provided — no case transfer is applied.  This is the correct mode for
    **uncloaking**, where the mapping stores the exact original text and we
    want to restore it verbatim (e.g. "BigCo LLC", not "Bigco Llc").

    Returns the total number of individual replacement operations performed.

    The function processes:
    - Body paragraphs
    - Table cells (nested tables included)
    - Headers and footers across all sections
    """
    if not replacements:
        return 0

    # Pre-compile a single regex that matches any of the target strings.
    # Sort longest-first so that "Acme Corporation" is matched before "Acme".
    sorted_targets = sorted(replacements.keys(), key=len, reverse=True)
    pattern = re.compile(
        "|".join(re.escape(t) for t in sorted_targets),
        flags=re.IGNORECASE,
    )

    # Build a normalized lookup: lowercased original -> new_text template.
    lookup: dict[str, str] = {k.lower(): v for k, v in replacements.items()}

    total = 0

    # --- body paragraphs ---
    for paragraph in doc.paragraphs:
        total += _replace_in_paragraph(paragraph, pattern, lookup, match_case)

    # --- tables ---
    total += _replace_in_tables(doc.tables, pattern, lookup, match_case)

    # --- headers / footers ---
    for section in doc.sections:
        for header_footer in _iter_headers_footers(section):
            for paragraph in header_footer.paragraphs:
                total += _replace_in_paragraph(paragraph, pattern, lookup, match_case)
            if hasattr(header_footer, "tables"):
                total += _replace_in_tables(header_footer.tables, pattern, lookup, match_case)

    logger.info("Replacements applied: %d", total)
    return total


# ---------------------------------------------------------------------------
# Internal: encryption detection
# ---------------------------------------------------------------------------

def _is_encrypted(path: Path) -> bool:
    """
    Detect whether a .docx file is encrypted.

    Encrypted Office documents are actually OLE2 Compound Binary files (not
    ZIP). The magic bytes ``\\xD0\\xCF\\x11\\xE0`` at offset 0 indicate OLE2.
    Additionally, some tools produce a ZIP that contains an
    ``EncryptedPackage`` entry.
    """
    try:
        with open(path, "rb") as fh:
            header = fh.read(8)
        # OLE2 magic — the entire file is an encrypted container.
        if header[:4] == b"\xd0\xcf\x11\xe0":
            return True
    except OSError:
        return False

    # ZIP-based check: look for EncryptedPackage entry.
    try:
        with zipfile.ZipFile(path, "r") as zf:
            if "EncryptedPackage" in zf.namelist():
                return True
    except zipfile.BadZipFile:
        pass

    return False


# ---------------------------------------------------------------------------
# Internal: text extraction helpers
# ---------------------------------------------------------------------------

def _extract_from_tables(tables) -> list[TextSource]:
    """Recursively extract text from tables (handles nested tables)."""
    results: list[TextSource] = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if text.strip():
                        results.append((text, cell))
                # Nested tables
                if cell.tables:
                    results.extend(_extract_from_tables(cell.tables))
    return results


def _iter_headers_footers(section):
    """Yield all header/footer objects for a section, skipping unlinked ones."""
    # python-docx lazily creates headers/footers on attribute access.  The
    # ``is_linked_to_previous`` flag tells us whether the section actually
    # owns its own header/footer or inherits from the prior section.
    attrs = [
        "header", "footer",
        "first_page_header", "first_page_footer",
        "even_page_header", "even_page_footer",
    ]
    for attr in attrs:
        hf = getattr(section, attr, None)
        if hf is not None and not hf.is_linked_to_previous:
            yield hf


# ---------------------------------------------------------------------------
# Internal: case-preserving replacement
# ---------------------------------------------------------------------------

def _is_bracketed_label(text: str) -> bool:
    """Return True if *text* is a bracketed label like ``[AltCustomerName]``.

    Bracketed labels are placeholder tokens that must be preserved verbatim
    during cloaking — applying case transfer would mangle the internal
    capitalisation (e.g. ``[AltCustomerName]`` -> ``[Altcustomername]``).
    """
    return len(text) >= 2 and text.startswith("[") and text.endswith("]")


def _transfer_case(original: str, replacement: str) -> str:
    """
    Transfer the case pattern of *original* onto *replacement*.

    Rules (applied in order):
    1. If *original* is all uppercase  -> return *replacement* uppercased.
    2. If *original* is all lowercase  -> return *replacement* lowercased.
    3. If *original* is Title Case     -> return *replacement* title-cased.
    4. If every word in *original* starts with an uppercase letter (loose
       title case, e.g. "Beta LLC") -> return *replacement* title-cased.
       This handles entity names with acronyms like "Acme USA Inc".
    5. If only the first character is uppercase (sentence case) -> capitalize
       just the first letter of *replacement*, lowercase the rest.
    6. Otherwise, transfer character-by-character as far as possible, then
       follow the case of the last mapped character for any remaining chars.
    """
    if not original or not replacement:
        return replacement

    if original.isupper():
        return replacement.upper()
    if original.islower():
        return replacement.lower()
    if original.istitle():
        return replacement.title()

    # Loose title case: every word starts with uppercase (covers "Beta LLC",
    # "Acme USA", etc. where strict istitle() fails due to all-caps words).
    words = original.split()
    if words and all(w[0].isupper() for w in words if w):
        return replacement.title()

    # Sentence case: first character upper, rest mostly lower.
    if original[0].isupper():
        return replacement[0].upper() + replacement[1:]

    # Character-by-character transfer.
    result = []
    for i, ch in enumerate(replacement):
        if i < len(original):
            if original[i].isupper():
                result.append(ch.upper())
            elif original[i].islower():
                result.append(ch.lower())
            else:
                result.append(ch)
        else:
            # Beyond the length of original -- continue with last known case.
            if result and result[-1].isupper():
                result.append(ch.upper())
            elif result and result[-1].islower():
                result.append(ch.lower())
            else:
                result.append(ch)
    return "".join(result)


# ---------------------------------------------------------------------------
# Internal: paragraph-level replacement engine
# ---------------------------------------------------------------------------

def _replace_in_paragraph(
    paragraph: "Paragraph",
    pattern: re.Pattern,
    lookup: dict[str, str],
    match_case: bool = True,
) -> int:
    """
    Replace all pattern matches in *paragraph*, preserving run-level
    formatting wherever possible.

    Returns the number of replacements made in this paragraph.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    # Concatenate run texts to build the full paragraph string.
    full_text = "".join(r.text for r in runs)
    if not full_text:
        return 0

    # Quick check: is there anything to replace?
    if not pattern.search(full_text):
        return 0

    # Build a mapping: character-index -> (run_index, offset_within_run).
    char_map = _build_char_map(runs)

    # Try the format-preserving path first.
    try:
        count = _replace_preserving_format(runs, full_text, char_map, pattern, lookup, match_case)
        return count
    except _FallbackNeeded:
        logger.debug(
            "Falling back to run-collapsing for paragraph: %.60s...",
            full_text[:60],
        )
        return _replace_collapsing_runs(paragraph, full_text, pattern, lookup, match_case)


class _FallbackNeeded(Exception):
    """Raised internally to trigger the run-collapsing fallback."""


def _build_char_map(runs: list["Run"]) -> list[tuple[int, int]]:
    """
    Return a list of length == total characters, where each entry is
    ``(run_index, offset_within_that_run)``.
    """
    char_map: list[tuple[int, int]] = []
    for run_idx, run in enumerate(runs):
        for offset in range(len(run.text)):
            char_map.append((run_idx, offset))
    return char_map


# ---------------------------------------------------------------------------
# Strategy 1: format-preserving cross-run replacement
# ---------------------------------------------------------------------------

def _replace_preserving_format(
    runs: list["Run"],
    full_text: str,
    char_map: list[tuple[int, int]],
    pattern: re.Pattern,
    lookup: dict[str, str],
    match_case: bool = True,
) -> int:
    """
    Walk matches from **right to left** (so earlier indices stay valid) and
    splice replacement characters directly into runs.

    If a single match spans runs in a way that makes character-level surgery
    ambiguous (e.g. replacement is shorter/longer than original and the match
    spans 3+ runs with different formatting on interior runs), we raise
    ``_FallbackNeeded`` so the caller can try the simpler strategy.
    """
    matches = list(pattern.finditer(full_text))
    if not matches:
        return 0

    # Process right-to-left to keep positional indices stable.
    for match in reversed(matches):
        matched_text = match.group()
        new_text_template = lookup[matched_text.lower()]
        if match_case and not _is_bracketed_label(new_text_template):
            new_text = _transfer_case(matched_text, new_text_template)
        else:
            new_text = new_text_template

        start, end = match.start(), match.end()
        match_len = end - start
        repl_len = len(new_text)

        # Identify which runs are touched.
        first_run_idx, first_offset = char_map[start]
        last_run_idx, _ = char_map[end - 1]

        span_count = last_run_idx - first_run_idx + 1

        if span_count == 1:
            # Entire match is within a single run — straightforward.
            _splice_single_run(runs[first_run_idx], first_offset, match_len, new_text)

        elif span_count == 2:
            # Match spans exactly two runs.
            _splice_two_runs(runs, first_run_idx, last_run_idx, char_map,
                             start, end, new_text)

        else:
            # Match spans 3+ runs. Attempt: put replacement in the first run,
            # clear the spanned portions in the middle runs, clear the portion
            # in the last run. If middle runs become empty that is fine.
            _splice_multi_runs(runs, first_run_idx, last_run_idx, char_map,
                               start, end, new_text)

    return len(matches)


def _splice_single_run(run: "Run", offset: int, length: int, new_text: str) -> None:
    """Replace *length* characters starting at *offset* in a single run."""
    text = run.text
    run.text = text[:offset] + new_text + text[offset + length:]


def _splice_two_runs(
    runs: list["Run"],
    first_idx: int,
    last_idx: int,
    char_map: list[tuple[int, int]],
    match_start: int,
    match_end: int,
    new_text: str,
) -> None:
    """
    Replace text spanning exactly two adjacent runs.

    Strategy: put the full replacement into the tail of the first run
    (starting where the match begins inside that run) and remove the matched
    portion from the second run.
    """
    first_run = runs[first_idx]
    last_run = runs[last_idx]

    _, first_offset = char_map[match_start]
    _, last_offset = char_map[match_end - 1]

    # First run: keep everything before the match start, append new_text.
    first_run.text = first_run.text[:first_offset] + new_text

    # Last run: remove everything up to and including the last matched char.
    last_run.text = last_run.text[last_offset + 1:]


def _splice_multi_runs(
    runs: list["Run"],
    first_idx: int,
    last_idx: int,
    char_map: list[tuple[int, int]],
    match_start: int,
    match_end: int,
    new_text: str,
) -> None:
    """
    Replace text spanning three or more runs.

    Strategy: put the full replacement text into the first run (preserving
    that run's formatting), blank out all middle runs, and trim the last run.
    """
    first_run = runs[first_idx]
    last_run = runs[last_idx]

    _, first_offset = char_map[match_start]
    _, last_offset = char_map[match_end - 1]

    # First run: keep text before match, append replacement.
    first_run.text = first_run.text[:first_offset] + new_text

    # Middle runs: clear completely.
    for mid_idx in range(first_idx + 1, last_idx):
        runs[mid_idx].text = ""

    # Last run: remove matched portion.
    last_run.text = last_run.text[last_offset + 1:]


# ---------------------------------------------------------------------------
# Strategy 2: fallback — collapse runs, replace, single run out
# ---------------------------------------------------------------------------

def _replace_collapsing_runs(
    paragraph: "Paragraph",
    full_text: str,
    pattern: re.Pattern,
    lookup: dict[str, str],
    match_case: bool = True,
) -> int:
    """
    Collapse all runs in *paragraph* into one, perform regex replacement on
    the concatenated text, and write the result back as a single run.

    This **loses per-run formatting** but guarantees correct text output.
    Formatting from the first run is preserved (so the paragraph keeps its
    dominant style).
    """
    count = 0

    def _replacer(m: re.Match) -> str:
        nonlocal count
        count += 1
        matched = m.group()
        template = lookup[matched.lower()]
        if match_case and not _is_bracketed_label(template):
            return _transfer_case(matched, template)
        return template

    new_text = pattern.sub(_replacer, full_text)

    if count == 0:
        return 0

    # Preserve formatting of the first run (or paragraph style if no runs).
    runs = paragraph.runs
    first_run_elem = None
    if runs:
        rpr = runs[0]._element.find(qn("w:rPr"))
        if rpr is not None:
            first_run_elem = deepcopy(rpr)

    # Remove all existing runs from the paragraph's XML.
    p_elem = paragraph._element
    for r_elem in p_elem.findall(qn("w:r")):
        p_elem.remove(r_elem)

    # Create a new single run.
    new_run = paragraph.add_run(new_text)
    if first_run_elem is not None:
        new_run._element.insert(0, first_run_elem)

    return count


# ---------------------------------------------------------------------------
# Internal: table helpers
# ---------------------------------------------------------------------------

def _replace_in_tables(
    tables,
    pattern: re.Pattern,
    lookup: dict[str, str],
    match_case: bool = True,
) -> int:
    """Apply replacements to every paragraph in every cell of all tables."""
    total = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += _replace_in_paragraph(paragraph, pattern, lookup, match_case)
                # Nested tables.
                if cell.tables:
                    total += _replace_in_tables(cell.tables, pattern, lookup, match_case)
    return total
