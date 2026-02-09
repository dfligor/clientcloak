"""
Document security scanning for ClientCloak.

Detects prompt injection attacks, hidden text, invisible Unicode characters,
and suspicious metadata in .docx documents before they are sent to AI services.

Prompt injection patterns are adapted from PlaybookRedliner's battle-tested
security module and extended with additional coverage.
"""

from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from .models import SecurityFinding, ThreatLevel


# ---------------------------------------------------------------------------
# Prompt injection patterns (from PlaybookRedliner + extensions)
# ---------------------------------------------------------------------------

INJECTION_PATTERNS: list[tuple[str, str]] = [
    # --- Direct instruction overrides (PlaybookRedliner core) ---
    (
        r"ignore\s+(all\s+)?(previous|prior|above|earlier)\s+(instructions?|rules?|guidelines?)",
        "Instruction override: ignore previous instructions",
    ),
    (
        r"disregard\s+(all\s+)?(previous|prior|above|earlier)\s+(instructions?|rules?|guidelines?)",
        "Instruction override: disregard previous instructions",
    ),
    (
        r"forget\s+(all\s+)?(previous|prior|above|earlier)\s+(instructions?|rules?|guidelines?)",
        "Instruction override: forget previous instructions",
    ),
    (
        r"override\s+(all\s+)?(previous|prior|above|earlier)\s+(instructions?|rules?|guidelines?)",
        "Instruction override: override previous instructions",
    ),
    # --- New instruction injections (PlaybookRedliner core) ---
    (
        r"new\s+instructions?:\s*",
        "Instruction injection: new instructions block",
    ),
    (
        r"system\s+prompt:\s*",
        "Instruction injection: system prompt reference",
    ),
    (
        r"admin\s+override:\s*",
        "Instruction injection: admin override",
    ),
    (
        r"developer\s+mode:\s*",
        "Instruction injection: developer mode",
    ),
    # --- Role manipulation (PlaybookRedliner core) ---
    (
        r"you\s+are\s+now\s+(a|an)\s+",
        "Role manipulation: 'you are now'",
    ),
    (
        r"pretend\s+(you\s+are|to\s+be)\s+",
        "Role manipulation: pretend directive",
    ),
    (
        r"act\s+as\s+(if\s+you\s+are\s+)?(a|an)\s+",
        "Role manipulation: act as directive",
    ),
    (
        r"roleplay\s+as\s+",
        "Role manipulation: roleplay directive",
    ),
    # --- Output manipulation (PlaybookRedliner core) ---
    (
        r"do\s+not\s+(flag|report|mention)\s+(this|the|any|issues?|problems?|findings?)",
        "Output manipulation: suppress findings",
    ),
    (
        r"skip\s+(all\s+)?rules?\s+",
        "Output manipulation: skip rules",
    ),
    (
        r"bypass\s+(all\s+)?rules?\s+",
        "Output manipulation: bypass rules",
    ),
    (
        r"return\s+(only\s+)?empty\s+(json|array|list)",
        "Output manipulation: force empty output",
    ),
    (
        r"output\s+(only\s+)?\[\s*\]",
        "Output manipulation: force empty array",
    ),
    (
        r"do\s+not\s+include\s+(this|in\s+(your|the)\s+(output|response|analysis|report))",
        "Output manipulation: suppress content in output",
    ),
    # --- Jailbreak attempts (PlaybookRedliner core) ---
    (
        r"DAN\s+mode",
        "Jailbreak: DAN mode reference",
    ),
    (
        r"jailbreak",
        "Jailbreak: explicit jailbreak reference",
    ),
    (
        r"ignore\s+safety",
        "Jailbreak: safety bypass",
    ),
    (
        r"ignore\s+ethical",
        "Jailbreak: ethics bypass",
    ),
    # --- Extended patterns (ClientCloak additions) ---
    (
        r"forget\s+everything",
        "Memory manipulation: forget everything",
    ),
    (
        r"clear\s+(your\s+)?(memory|context|history)",
        "Memory manipulation: clear context",
    ),
    (
        r"reset\s+(your\s+)?(memory|context|instructions?)",
        "Memory manipulation: reset context",
    ),
    (
        r"disregard\s+(the\s+)?above",
        "Instruction override: disregard the above",
    ),
    (
        r"from\s+now\s+on,?\s+(you|ignore|only|always|never)",
        "Instruction injection: 'from now on' directive",
    ),
    (
        r"begin\s+(new\s+)?session",
        "Context manipulation: session reset",
    ),
    (
        r"end\s+(of\s+)?system\s+prompt",
        "Boundary attack: end of system prompt",
    ),
    (
        r"\[INST\]|\[/INST\]|<<SYS>>|<</SYS>>",
        "Boundary attack: LLM control tokens",
    ),
    (
        r"<\|im_start\|>|<\|im_end\|>",
        "Boundary attack: ChatML control tokens",
    ),
    (
        r"human:\s*\n|assistant:\s*\n",
        "Boundary attack: conversation role injection",
    ),
    (
        r"IMPORTANT:\s*(ignore|disregard|forget|override|bypass)",
        "Social engineering: urgency-prefixed override",
    ),
    (
        r"do\s+not\s+reveal\s+(your|the|any)\s+(system|instructions?|prompt|rules?)",
        "Exfiltration attempt: suppress system prompt disclosure",
    ),
    (
        r"repeat\s+(your|the)\s+(system|initial)\s+(prompt|instructions?|message)",
        "Exfiltration attempt: system prompt extraction",
    ),
    (
        r"what\s+(are|were)\s+your\s+(initial|system|original)\s+(instructions?|prompt|rules?)",
        "Exfiltration attempt: system prompt query",
    ),
]

# Compile all patterns for efficient reuse across scans.
COMPILED_INJECTION_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(pattern, re.IGNORECASE | re.MULTILINE), description)
    for pattern, description in INJECTION_PATTERNS
]


# ---------------------------------------------------------------------------
# Invisible Unicode character map
# ---------------------------------------------------------------------------

INVISIBLE_CHARS: dict[int, str] = {
    0x200B: "Zero-Width Space",
    0x200C: "Zero-Width Non-Joiner",
    0x200D: "Zero-Width Joiner",
    0x200E: "Left-to-Right Mark",
    0x200F: "Right-to-Left Mark",
    0x202A: "Left-to-Right Embedding",
    0x202B: "Right-to-Left Embedding",
    0x202C: "Pop Directional Formatting",
    0x202D: "Left-to-Right Override",
    0x202E: "Right-to-Left Override",
    0x2060: "Word Joiner",
    0x2061: "Function Application",
    0x2062: "Invisible Times",
    0x2063: "Invisible Separator",
    0x2064: "Invisible Plus",
    0x2066: "Left-to-Right Isolate",
    0x2067: "Right-to-Left Isolate",
    0x2068: "First Strong Isolate",
    0x2069: "Pop Directional Isolate",
    0x00AD: "Soft Hyphen",
    0xFEFF: "Byte Order Mark / Zero-Width No-Break Space",
    0xFFF9: "Interlinear Annotation Anchor",
    0xFFFA: "Interlinear Annotation Separator",
    0xFFFB: "Interlinear Annotation Terminator",
    0x180E: "Mongolian Vowel Separator",
    0x034F: "Combining Grapheme Joiner",
}

# Pre-build a single character class regex for efficient scanning.
_INVISIBLE_CHAR_CLASS = "[" + "".join(
    rf"\u{cp:04X}" for cp in sorted(INVISIBLE_CHARS.keys())
) + "]"
INVISIBLE_CHAR_RE = re.compile(_INVISIBLE_CHAR_CLASS)


# ---------------------------------------------------------------------------
# Color helpers (adapted from PlaybookRedliner)
# ---------------------------------------------------------------------------

def _is_light_color(color: RGBColor | None, threshold: int = 240) -> bool:
    """
    Return True if *color* is near-white (all channels >= *threshold*).

    Mirrors PlaybookRedliner's ``is_light_color`` logic.
    """
    if color is None:
        return False
    try:
        return color[0] >= threshold and color[1] >= threshold and color[2] >= threshold
    except (TypeError, AttributeError, IndexError):
        return False


def _color_hex(color: RGBColor | None) -> str:
    """Return a human-readable hex string for an RGBColor, or 'N/A'."""
    if color is None:
        return "N/A"
    try:
        return f"#{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    except (TypeError, AttributeError, IndexError):
        return "N/A"


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_all_text_for_scanning(doc: Document) -> str:
    """
    Extract every text surface from *doc* for pattern-matching.

    Sources: paragraphs, tables, headers, footers.  Returns a single
    string with newlines between logical blocks.
    """
    parts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    # Headers and footers across all sections
    for section in doc.sections:
        for header_footer in (section.header, section.first_page_header,
                              section.even_page_header, section.footer,
                              section.first_page_footer, section.even_page_footer):
            # Even "linked to previous" headers inherit content from the
            # prior section, so we scan all of them unconditionally.
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Individual detectors
# ---------------------------------------------------------------------------

def detect_hidden_text(doc: Document) -> list[SecurityFinding]:
    """
    Detect text hidden via tiny fonts, the hidden font attribute, or
    near-white coloring.

    Returns a list of SecurityFinding instances (may be empty).
    """
    findings: list[SecurityFinding] = []

    def _scan_runs(runs, location_prefix: str) -> None:
        for run in runs:
            run_text = run.text.strip()
            if not run_text:
                continue

            font = run.font

            # --- Tiny font (< 2pt) ---
            if font.size is not None:
                pt_size = font.size / 12700  # EMU -> points
                if pt_size < 2:
                    findings.append(SecurityFinding(
                        threat_level=ThreatLevel.CRITICAL,
                        finding_type="hidden_text",
                        description=(
                            f"Text rendered at {pt_size:.1f}pt — effectively invisible to readers."
                        ),
                        location=location_prefix,
                        content_preview=run_text[:200],
                        recommendation="Remove the hidden text or increase the font size.",
                    ))

            # --- Hidden font attribute ---
            # python-docx exposes font.hidden but the underlying XML attribute
            # may also be set directly.  Check both.
            hidden = font.hidden
            if hidden is None:
                # Fall back to raw XML check
                rpr = run._element.find(qn("w:rPr"))
                if rpr is not None and rpr.find(qn("w:vanish")) is not None:
                    hidden = True
            if hidden:
                findings.append(SecurityFinding(
                    threat_level=ThreatLevel.CRITICAL,
                    finding_type="hidden_text",
                    description="Text has the 'hidden' font attribute enabled.",
                    location=location_prefix,
                    content_preview=run_text[:200],
                    recommendation="Remove the hidden text or clear the hidden attribute.",
                ))

            # --- Near-white color ---
            if font.color and font.color.rgb:
                if _is_light_color(font.color.rgb):
                    hex_val = _color_hex(font.color.rgb)
                    findings.append(SecurityFinding(
                        threat_level=ThreatLevel.CRITICAL,
                        finding_type="hidden_text",
                        description=(
                            f"Text color is near-white ({hex_val}), making it invisible "
                            f"on a white background."
                        ),
                        location=location_prefix,
                        content_preview=run_text[:200],
                        recommendation="Remove the hidden text or change the font color.",
                    ))

    # Paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        _scan_runs(para.runs, f"Paragraph {para_idx + 1}")

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    loc = (
                        f"Table {table_idx + 1}, "
                        f"Row {row_idx + 1}, "
                        f"Cell {cell_idx + 1}"
                    )
                    _scan_runs(para.runs, loc)

    # Headers / footers
    for sec_idx, section in enumerate(doc.sections):
        for label, hf in (
            ("Header", section.header),
            ("First Page Header", section.first_page_header),
            ("Even Page Header", section.even_page_header),
            ("Footer", section.footer),
            ("First Page Footer", section.first_page_footer),
            ("Even Page Footer", section.even_page_footer),
        ):
            if hf is None:
                continue
            for para in hf.paragraphs:
                loc = f"Section {sec_idx + 1} {label}"
                _scan_runs(para.runs, loc)

    return findings


def detect_prompt_injection(doc: Document) -> list[SecurityFinding]:
    """
    Scan all document text surfaces for prompt injection patterns.

    Uses the compiled pattern set originally from PlaybookRedliner,
    extended with additional coverage for ClientCloak.
    """
    findings: list[SecurityFinding] = []
    seen: set[tuple[str, str]] = set()  # (pattern_desc, location) dedup

    def _check_text(text: str, location: str) -> None:
        if not text:
            return
        for regex, description in COMPILED_INJECTION_PATTERNS:
            match = regex.search(text)
            if match:
                key = (description, location)
                if key in seen:
                    continue
                seen.add(key)

                # Build a snippet centred around the match
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + 30)
                snippet = text[start:end]
                if start > 0:
                    snippet = "..." + snippet
                if end < len(text):
                    snippet = snippet + "..."

                findings.append(SecurityFinding(
                    threat_level=ThreatLevel.CRITICAL,
                    finding_type="prompt_injection",
                    description=description,
                    location=location,
                    content_preview=snippet[:200],
                    recommendation=(
                        "Remove or rewrite this text before sending the document to "
                        "an AI service."
                    ),
                ))

    # Body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        _check_text(para.text, f"Paragraph {para_idx + 1}")

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                loc = (
                    f"Table {table_idx + 1}, "
                    f"Row {row_idx + 1}, "
                    f"Cell {cell_idx + 1}"
                )
                _check_text(cell.text, loc)

    # Headers / footers
    for sec_idx, section in enumerate(doc.sections):
        for label, hf in (
            ("Header", section.header),
            ("First Page Header", section.first_page_header),
            ("Even Page Header", section.even_page_header),
            ("Footer", section.footer),
            ("First Page Footer", section.first_page_footer),
            ("Even Page Footer", section.even_page_footer),
        ):
            if hf is None:
                continue
            for para in hf.paragraphs:
                loc = f"Section {sec_idx + 1} {label}"
                _check_text(para.text, loc)

    return findings


def detect_invisible_characters(doc: Document) -> list[SecurityFinding]:
    """
    Detect invisible Unicode characters that could carry hidden payloads
    or manipulate text rendering.
    """
    findings: list[SecurityFinding] = []
    seen_locations: set[str] = set()

    def _check_text(text: str, location: str) -> None:
        if not text:
            return
        chars_found: dict[str, int] = {}
        for match in INVISIBLE_CHAR_RE.finditer(text):
            cp = ord(match.group())
            name = INVISIBLE_CHARS.get(cp, f"U+{cp:04X}")
            chars_found[name] = chars_found.get(name, 0) + 1

        if chars_found:
            key = location
            if key in seen_locations:
                return
            seen_locations.add(key)

            detail_parts = [f"{name} x{count}" for name, count in chars_found.items()]
            total = sum(chars_found.values())

            # Determine threat level based on count
            threat = ThreatLevel.WARNING if total < 5 else ThreatLevel.CRITICAL

            # Build a preview showing the surrounding visible text
            visible = INVISIBLE_CHAR_RE.sub("\u2423", text)  # replace with open box
            preview = visible[:200]

            findings.append(SecurityFinding(
                threat_level=threat,
                finding_type="invisible_chars",
                description=(
                    f"Found {total} invisible character(s): {', '.join(detail_parts)}."
                ),
                location=location,
                content_preview=preview,
                recommendation=(
                    "Remove invisible characters. They may carry hidden payloads or "
                    "manipulate text rendering/AI interpretation."
                ),
            ))

    # Body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        _check_text(para.text, f"Paragraph {para_idx + 1}")

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                loc = (
                    f"Table {table_idx + 1}, "
                    f"Row {row_idx + 1}, "
                    f"Cell {cell_idx + 1}"
                )
                _check_text(cell.text, loc)

    # Headers / footers
    for sec_idx, section in enumerate(doc.sections):
        for label, hf in (
            ("Header", section.header),
            ("First Page Header", section.first_page_header),
            ("Even Page Header", section.even_page_header),
            ("Footer", section.footer),
            ("First Page Footer", section.first_page_footer),
            ("Even Page Footer", section.even_page_footer),
        ):
            if hf is None:
                continue
            for para in hf.paragraphs:
                loc = f"Section {sec_idx + 1} {label}"
                _check_text(para.text, loc)

    return findings


def scan_metadata_fields(doc: Document) -> list[SecurityFinding]:
    """
    Check document core properties for suspicious content such as
    prompt injection patterns embedded in metadata fields.
    """
    findings: list[SecurityFinding] = []
    props = doc.core_properties

    fields: list[tuple[str, str | None]] = [
        ("Author", props.author),
        ("Last Modified By", props.last_modified_by),
        ("Subject", props.subject),
        ("Keywords", props.keywords),
        ("Comments", props.comments),
        ("Title", props.title),
        ("Category", props.category),
    ]

    for field_name, value in fields:
        if not value:
            continue

        # Check for injection patterns inside metadata
        for regex, description in COMPILED_INJECTION_PATTERNS:
            if regex.search(value):
                findings.append(SecurityFinding(
                    threat_level=ThreatLevel.CRITICAL,
                    finding_type="metadata",
                    description=(
                        f"Prompt injection pattern in '{field_name}' metadata: "
                        f"{description}"
                    ),
                    location=f"Document Properties > {field_name}",
                    content_preview=value[:200],
                    recommendation=(
                        f"Clear or rewrite the '{field_name}' metadata field before "
                        f"sending this document to an AI service."
                    ),
                ))
                break  # one finding per field is sufficient

        # Check for invisible characters in metadata
        invisible_count = len(INVISIBLE_CHAR_RE.findall(value))
        if invisible_count:
            findings.append(SecurityFinding(
                threat_level=ThreatLevel.WARNING,
                finding_type="metadata",
                description=(
                    f"Found {invisible_count} invisible character(s) in "
                    f"'{field_name}' metadata."
                ),
                location=f"Document Properties > {field_name}",
                content_preview=value[:200],
                recommendation=(
                    f"Clean the '{field_name}' metadata field to remove invisible "
                    f"characters."
                ),
            ))

    return findings


# ---------------------------------------------------------------------------
# Finding removal
# ---------------------------------------------------------------------------

def remove_finding(doc: Document, finding: SecurityFinding) -> bool:
    """
    Attempt to remove or neutralise the content identified by *finding*.

    Returns True if the content was successfully removed/neutralised,
    False if the finding type or location could not be resolved.

    Supported finding types:
    - hidden_text: removes the offending run entirely.
    - prompt_injection: clears text of the paragraph/cell containing the match.
    - invisible_chars: strips invisible characters in-place.
    - metadata: clears the relevant metadata field.
    """
    if finding.finding_type == "metadata":
        return _remove_metadata_finding(doc, finding)

    if finding.finding_type in ("hidden_text", "prompt_injection", "invisible_chars"):
        return _remove_content_finding(doc, finding)

    return False


def _remove_metadata_finding(doc: Document, finding: SecurityFinding) -> bool:
    """Clear a metadata field identified in the finding location."""
    # Location format: "Document Properties > FieldName"
    props = doc.core_properties
    field_map: dict[str, str] = {
        "Author": "author",
        "Last Modified By": "last_modified_by",
        "Subject": "subject",
        "Keywords": "keywords",
        "Comments": "comments",
        "Title": "title",
        "Category": "category",
    }
    for display_name, attr_name in field_map.items():
        if display_name in finding.location:
            try:
                setattr(props, attr_name, "")
                return True
            except (AttributeError, TypeError):
                return False
    return False


def _remove_content_finding(doc: Document, finding: SecurityFinding) -> bool:
    """
    Remove or neutralise content from the document body based on a finding's
    location string and type.
    """
    location = finding.location

    # --- Paragraph-level findings ---
    para_match = re.match(r"Paragraph (\d+)", location)
    if para_match:
        para_idx = int(para_match.group(1)) - 1
        if 0 <= para_idx < len(doc.paragraphs):
            return _apply_removal(doc.paragraphs[para_idx], finding)
        return False

    # --- Table cell findings ---
    table_match = re.match(
        r"Table (\d+), Row (\d+), Cell (\d+)", location
    )
    if table_match:
        t_idx = int(table_match.group(1)) - 1
        r_idx = int(table_match.group(2)) - 1
        c_idx = int(table_match.group(3)) - 1
        try:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
            removed = False
            for para in cell.paragraphs:
                if _apply_removal(para, finding):
                    removed = True
            return removed
        except (IndexError, KeyError):
            return False

    # --- Header / footer findings ---
    section_match = re.match(r"Section (\d+) (.+)", location)
    if section_match:
        sec_idx = int(section_match.group(1)) - 1
        hf_label = section_match.group(2)
        if 0 <= sec_idx < len(doc.sections):
            section = doc.sections[sec_idx]
            hf_map: dict[str, object] = {
                "Header": section.header,
                "First Page Header": section.first_page_header,
                "Even Page Header": section.even_page_header,
                "Footer": section.footer,
                "First Page Footer": section.first_page_footer,
                "Even Page Footer": section.even_page_footer,
            }
            hf = hf_map.get(hf_label)
            if hf is not None:
                removed = False
                for para in hf.paragraphs:
                    if _apply_removal(para, finding):
                        removed = True
                return removed
        return False

    return False


def _apply_removal(paragraph, finding: SecurityFinding) -> bool:
    """
    Apply the appropriate removal strategy to a paragraph based on
    finding type.
    """
    if finding.finding_type == "hidden_text":
        # Remove runs that match the hidden-text criteria
        removed = False
        for run in list(paragraph.runs):
            run_text = run.text.strip()
            if not run_text:
                continue
            # Check if this run matches the finding's content preview
            if finding.content_preview and run_text[:50] in finding.content_preview:
                run._element.getparent().remove(run._element)
                removed = True
                continue
            # Also remove any run that exhibits hidden-text traits
            font = run.font
            is_hidden = (
                (font.size is not None and font.size / 12700 < 2)
                or font.hidden is True
                or (font.color and font.color.rgb and _is_light_color(font.color.rgb))
            )
            if is_hidden and run_text[:50] in finding.content_preview:
                run._element.getparent().remove(run._element)
                removed = True
        return removed

    if finding.finding_type == "prompt_injection":
        # Clear the paragraph text entirely (the whole paragraph is suspect)
        if finding.content_preview:
            # Verify this is the right paragraph by checking overlap
            para_text = paragraph.text
            # Use the content_preview (minus ellipsis) to confirm
            preview_core = finding.content_preview.strip(".")
            if preview_core in para_text or para_text[:50] in finding.content_preview:
                for run in paragraph.runs:
                    run.text = ""
                return True
        return False

    if finding.finding_type == "invisible_chars":
        # Strip invisible characters from every run in the paragraph
        modified = False
        for run in paragraph.runs:
            cleaned = INVISIBLE_CHAR_RE.sub("", run.text)
            if cleaned != run.text:
                run.text = cleaned
                modified = True
        return modified

    return False


# ---------------------------------------------------------------------------
# Master scan
# ---------------------------------------------------------------------------

_THREAT_SORT_ORDER = {
    ThreatLevel.CRITICAL: 0,
    ThreatLevel.WARNING: 1,
    ThreatLevel.INFO: 2,
}


def scan_document(doc: Document) -> list[SecurityFinding]:
    """
    Run all security detectors against *doc* and return findings sorted
    by threat level (critical first).

    This is the primary entry point for the ClientCloak security scanner.
    """
    findings: list[SecurityFinding] = []

    findings.extend(detect_hidden_text(doc))
    findings.extend(detect_prompt_injection(doc))
    findings.extend(detect_invisible_characters(doc))
    findings.extend(scan_metadata_fields(doc))

    # Stable sort: critical -> warning -> info
    findings.sort(key=lambda f: _THREAT_SORT_ORDER.get(f.threat_level, 99))

    return findings
