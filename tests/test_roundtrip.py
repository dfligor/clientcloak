"""
End-to-end roundtrip tests: create doc -> cloak -> verify sanitized -> uncloak -> verify restored.

Key test: case preservation for entity names like "BigCo LLC" through the
cloak/uncloak cycle.
"""

import pytest
from pathlib import Path
from docx import Document

from clientcloak.cloaker import cloak_document, sanitize_filename, sanitize_filename_for_config, _build_cloak_replacements, _expand_content_replacements
from clientcloak.uncloaker import uncloak_document
from clientcloak.docx_handler import load_document, extract_all_text
from clientcloak.models import CloakConfig, CommentMode, PartyAlias
from tests.conftest import make_simple_docx


# ===================================================================
# Roundtrip tests
# ===================================================================

class TestRoundtrip:
    """Full roundtrip: create -> cloak -> uncloak -> verify."""

    def _do_roundtrip(self, tmp_path, paragraphs, party_a, party_b,
                       label_a="Licensor", label_b="Licensee",
                       comment_mode=CommentMode.STRIP):
        """Helper: create a doc, cloak it, uncloak it, return original and final texts."""
        # Create the document
        input_path = make_simple_docx(tmp_path / "original.docx", paragraphs)
        cloaked_path = tmp_path / "cloaked.docx"
        mapping_path = tmp_path / "mapping.json"
        uncloaked_path = tmp_path / "uncloaked.docx"

        config = CloakConfig(
            party_a_name=party_a,
            party_a_label=label_a,
            party_b_name=party_b,
            party_b_label=label_b,
            comment_mode=comment_mode,
            strip_metadata=True,
        )

        # Cloak
        result = cloak_document(input_path, cloaked_path, mapping_path, config)
        assert result.replacements_applied > 0

        # Uncloak
        uncloak_count = uncloak_document(cloaked_path, uncloaked_path, mapping_path)
        assert uncloak_count > 0

        # Read texts
        original_texts = [p.text for p in load_document(input_path).paragraphs if p.text.strip()]
        cloaked_texts = [p.text for p in load_document(cloaked_path).paragraphs if p.text.strip()]
        uncloaked_texts = [p.text for p in load_document(uncloaked_path).paragraphs if p.text.strip()]

        return original_texts, cloaked_texts, uncloaked_texts

    def test_basic_roundtrip(self, tmp_path):
        """Basic cloak/uncloak cycle restores original text."""
        paragraphs = [
            "This agreement is between Acme Corporation and BigCo LLC.",
            "Acme Corporation shall deliver the goods.",
            "BigCo LLC shall make payment.",
        ]
        original, cloaked, uncloaked = self._do_roundtrip(
            tmp_path, paragraphs, "Acme Corporation", "BigCo LLC"
        )

        # Cloaked should not contain original names
        for text in cloaked:
            assert "Acme Corporation" not in text
            assert "BigCo LLC" not in text

        # Cloaked should contain labels
        cloaked_full = " ".join(cloaked)
        assert "Licensor" in cloaked_full
        assert "Licensee" in cloaked_full

        # Uncloaked should restore original names
        for orig, final in zip(original, uncloaked):
            assert orig == final, f"Roundtrip mismatch: {orig!r} != {final!r}"

    def test_case_preservation_bigco_llc(self, tmp_path):
        """
        Critical test: 'BigCo LLC' should round-trip correctly.

        During cloaking, 'BigCo LLC' -> 'Licensee' (case transfer applies).
        During uncloaking with match_case=False, 'Licensee' -> 'BigCo LLC' verbatim.
        The result should be 'BigCo LLC', NOT 'Bigco Llc'.
        """
        paragraphs = [
            "BigCo LLC is the customer.",
            "The obligations of BigCo LLC include payment.",
        ]
        original, cloaked, uncloaked = self._do_roundtrip(
            tmp_path, paragraphs, "Acme Corp", "BigCo LLC",
            label_a="Vendor", label_b="Customer",
        )

        for text in uncloaked:
            if "BigCo LLC" in original[uncloaked.index(text)]:
                assert "BigCo LLC" in text, (
                    f"Case not preserved: expected 'BigCo LLC' in '{text}'"
                )

    def test_all_caps_roundtrip(self, tmp_path):
        """
        Case handling through cloak/uncloak cycle:

        Cloaking (match_case=True): bracketed labels are preserved verbatim
          'ACME CORPORATION' -> '[Licensor]' (no case transfer for bracketed labels)
          'Acme Corporation' -> '[Licensor]' (verbatim)
          'acme corporation' -> '[Licensor]' (verbatim)

        Uncloaking (match_case=False): replacement is verbatim from mapping
          The mapping stores: '[Licensor]' -> 'Acme Corporation'
          So all instances of '[Licensor]' get replaced with 'Acme Corporation'.
        """
        paragraphs = [
            "ACME CORPORATION agrees to the terms.",
            "Acme Corporation is the vendor.",
            "acme corporation shall comply.",
        ]
        _, cloaked, uncloaked = self._do_roundtrip(
            tmp_path, paragraphs, "Acme Corporation", "BigCo LLC"
        )

        # Cloaking phase: bracketed labels are used verbatim regardless of source case
        assert "[Licensor]" in cloaked[0]
        assert "[Licensor]" in cloaked[1]
        assert "[Licensor]" in cloaked[2]

        # Uncloaking with match_case=False restores verbatim from mapping
        # All instances of '[Licensor]' become 'Acme Corporation'
        for text in uncloaked:
            assert "Acme Corporation" in text

    def test_roundtrip_with_table(self, tmp_path):
        """Roundtrip works for text inside table cells."""
        from tests.conftest import make_table_docx

        input_path = make_table_docx(
            tmp_path / "table_original.docx",
            [
                ["Party", "Role"],
                ["Acme Corporation", "Vendor"],
                ["BigCo LLC", "Customer"],
            ],
        )
        cloaked_path = tmp_path / "table_cloaked.docx"
        mapping_path = tmp_path / "table_mapping.json"
        uncloaked_path = tmp_path / "table_uncloaked.docx"

        config = CloakConfig(
            party_a_name="Acme Corporation",
            party_a_label="Licensor",
            party_b_name="BigCo LLC",
            party_b_label="Licensee",
            strip_metadata=True,
        )

        cloak_document(input_path, cloaked_path, mapping_path, config)

        # Verify cloaked table
        cloaked_doc = load_document(cloaked_path)
        cloaked_texts = extract_all_text(cloaked_doc)
        cloaked_text_strs = [t for t, _ in cloaked_texts]
        assert not any("Acme Corporation" in t for t in cloaked_text_strs)

        # Uncloak
        uncloak_document(cloaked_path, uncloaked_path, mapping_path)
        uncloaked_doc = load_document(uncloaked_path)
        uncloaked_texts = extract_all_text(uncloaked_doc)
        uncloaked_text_strs = [t for t, _ in uncloaked_texts]
        assert any("Acme Corporation" in t for t in uncloaked_text_strs)
        assert any("BigCo LLC" in t for t in uncloaked_text_strs)

    def test_roundtrip_preserves_document_structure(self, tmp_path):
        """Verify that paragraph count is preserved through roundtrip."""
        paragraphs = [
            "First paragraph with Acme Corporation.",
            "Second paragraph with BigCo LLC.",
            "Third paragraph without names.",
        ]
        input_path = make_simple_docx(tmp_path / "struct.docx", paragraphs)
        cloaked_path = tmp_path / "struct_cloaked.docx"
        mapping_path = tmp_path / "struct_mapping.json"

        config = CloakConfig(
            party_a_name="Acme Corporation",
            party_a_label="Licensor",
            party_b_name="BigCo LLC",
            party_b_label="Licensee",
        )
        cloak_document(input_path, cloaked_path, mapping_path, config)

        original_doc = load_document(input_path)
        cloaked_doc = load_document(cloaked_path)
        # Non-empty paragraph count should be the same
        orig_count = sum(1 for p in original_doc.paragraphs if p.text.strip())
        cloaked_count = sum(1 for p in cloaked_doc.paragraphs if p.text.strip())
        assert orig_count == cloaked_count


# ===================================================================
# Alias roundtrip tests
# ===================================================================

class TestAliasRoundtrip:
    """Verify that party aliases cloak and uncloak correctly."""

    def test_alias_basic_roundtrip(self, tmp_path):
        """Cloak+uncloak with aliases: all names replaced and restored."""
        paragraphs = [
            'This agreement is between Acme Corp. ("Acme") and BigCo LLC.',
            "Acme Corp. shall deliver the goods.",
            "Acme shall comply with all regulations.",
        ]
        input_path = make_simple_docx(tmp_path / "alias.docx", paragraphs)
        cloaked_path = tmp_path / "alias_cloaked.docx"
        mapping_path = tmp_path / "alias_mapping.json"
        uncloaked_path = tmp_path / "alias_uncloaked.docx"

        config = CloakConfig(
            party_a_name="Acme Corp.",
            party_a_label="Full Vendor Name",
            party_b_name="BigCo LLC",
            party_b_label="Customer",
            party_a_aliases=[PartyAlias(name="Acme", label="Vendor")],
        )

        result = cloak_document(input_path, cloaked_path, mapping_path, config)
        assert result.replacements_applied > 0

        # Verify cloaked text has no original names
        cloaked_doc = load_document(cloaked_path)
        cloaked_text = " ".join(p.text for p in cloaked_doc.paragraphs)
        assert "Acme Corp." not in cloaked_text
        assert "BigCo LLC" not in cloaked_text
        # The standalone "Acme" should also be replaced
        assert "[Vendor]" in cloaked_text
        assert "[Full Vendor Name]" in cloaked_text

        # Uncloak and verify restoration
        uncloak_document(cloaked_path, uncloaked_path, mapping_path)
        uncloaked_doc = load_document(uncloaked_path)
        for orig_p, final_p in zip(
            load_document(input_path).paragraphs, uncloaked_doc.paragraphs
        ):
            assert orig_p.text == final_p.text

    def test_alias_longest_first_ordering(self, tmp_path):
        """'Acme Corporation' matches before 'Acme' (no double-replacement)."""
        paragraphs = [
            "Acme Corporation is the vendor.",
            "Acme is the short name.",
        ]
        input_path = make_simple_docx(tmp_path / "order.docx", paragraphs)
        cloaked_path = tmp_path / "order_cloaked.docx"
        mapping_path = tmp_path / "order_mapping.json"

        config = CloakConfig(
            party_a_name="Acme Corporation",
            party_a_label="Full Vendor",
            party_b_name="BigCo LLC",
            party_b_label="Customer",
            party_a_aliases=[PartyAlias(name="Acme", label="Vendor")],
        )

        cloak_document(input_path, cloaked_path, mapping_path, config)
        cloaked_doc = load_document(cloaked_path)
        texts = [p.text for p in cloaked_doc.paragraphs if p.text.strip()]

        # "Acme Corporation" should become "[Full Vendor]", not "[Vendor] Corporation"
        assert "[Full Vendor]" in texts[0]
        assert "Corporation" not in texts[0]
        # Standalone "Acme" should become "[Vendor]"
        assert "[Vendor]" in texts[1]

    def test_no_aliases_backward_compatible(self, tmp_path):
        """CloakConfig with no aliases works identically to before."""
        paragraphs = [
            "Acme Corporation and BigCo LLC agree to the terms.",
        ]
        input_path = make_simple_docx(tmp_path / "compat.docx", paragraphs)
        cloaked_path = tmp_path / "compat_cloaked.docx"
        mapping_path = tmp_path / "compat_mapping.json"
        uncloaked_path = tmp_path / "compat_uncloaked.docx"

        config = CloakConfig(
            party_a_name="Acme Corporation",
            party_a_label="Licensor",
            party_b_name="BigCo LLC",
            party_b_label="Licensee",
        )

        result = cloak_document(input_path, cloaked_path, mapping_path, config)
        assert result.replacements_applied > 0
        assert config.party_a_aliases == []
        assert config.party_b_aliases == []

        uncloak_document(cloaked_path, uncloaked_path, mapping_path)
        orig_text = load_document(input_path).paragraphs[0].text
        final_text = load_document(uncloaked_path).paragraphs[0].text
        assert orig_text == final_text


# ===================================================================
# Filename sanitization tests
# ===================================================================

class TestFilenameSanitization:
    """Verify that party names are replaced in output filenames."""

    def test_sanitize_filename_basic(self):
        """Party names in the filename stem are replaced by bracketed labels."""
        replacements = {"Acme": "[Customer]", "BigCo": "[Vendor]"}
        result = sanitize_filename("Acme_BigCo_NDA", replacements)
        assert result == "[Customer]_[Vendor]_NDA"

    def test_sanitize_filename_case_insensitive(self):
        """Replacement is case-insensitive on the filename."""
        replacements = {"Acme": "[Customer]", "BigCo": "[Vendor]"}
        result = sanitize_filename("acme_bigco_NDA", replacements)
        assert result == "[Customer]_[Vendor]_NDA"

    def test_sanitize_filename_mixed_case(self):
        """Mixed-case party names in the filename are still replaced."""
        replacements = {"Acme": "[Customer]", "BigCo": "[Vendor]"}
        result = sanitize_filename("ACME_BIGCO_NDA", replacements)
        assert result == "[Customer]_[Vendor]_NDA"

    def test_sanitize_filename_no_match(self):
        """A filename with no party names is returned unchanged."""
        replacements = {"Acme": "[Customer]", "BigCo": "[Vendor]"}
        result = sanitize_filename("contract_NDA", replacements)
        assert result == "contract_NDA"

    def test_sanitize_filename_longest_first(self):
        """Longer original names are replaced before shorter ones."""
        replacements = {
            "Acme Corporation": "[Full Vendor]",
            "Acme": "[Vendor]",
        }
        result = sanitize_filename("Acme_Corporation_contract", replacements)
        assert result == "[Full Vendor]_contract"

    def test_sanitize_filename_for_config(self):
        """The config-based convenience wrapper works correctly."""
        config = CloakConfig(
            party_a_name="Acme",
            party_a_label="Customer",
            party_b_name="BigCo",
            party_b_label="Vendor",
        )
        result = sanitize_filename_for_config("Acme_BigCo_NDA", config)
        assert result == "[Customer]_[Vendor]_NDA"

    def test_cloak_document_sanitizes_output_filename(self, tmp_path):
        """
        End-to-end: cloak_document rewrites the output filename so party
        names do not leak.  Input 'Acme_BigCo_NDA.docx' with party_a='Acme'
        and party_b='BigCo' should produce '[Customer]_[Vendor]_NDA_cloaked.docx'.
        """
        paragraphs = ["This agreement is between Acme and BigCo."]
        input_path = make_simple_docx(tmp_path / "Acme_BigCo_NDA.docx", paragraphs)

        # The caller requests this output name (which still contains party names):
        requested_output = tmp_path / "Acme_BigCo_NDA_cloaked.docx"
        mapping_path = tmp_path / "mapping.json"

        config = CloakConfig(
            party_a_name="Acme",
            party_a_label="Customer",
            party_b_name="BigCo",
            party_b_label="Vendor",
        )

        result = cloak_document(input_path, requested_output, mapping_path, config)

        # The actual output path should have sanitized the filename.
        actual_output = Path(result.output_path)
        assert actual_output.name == "[Customer]_[Vendor]_NDA_cloaked.docx"
        assert actual_output.exists()

        # The original (unsanitized) path should NOT exist.
        assert not requested_output.exists()

        # The content should also be sanitized.
        doc = load_document(actual_output)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Acme" not in full_text
        assert "BigCo" not in full_text

    def test_cloak_document_filename_case_insensitive(self, tmp_path):
        """
        Output filename sanitization is case-insensitive: 'acme' in the
        filename matches party_a_name='Acme'.
        """
        paragraphs = ["Agreement between Acme and BigCo."]
        input_path = make_simple_docx(tmp_path / "acme_bigco_nda.docx", paragraphs)

        requested_output = tmp_path / "acme_bigco_nda_cloaked.docx"
        mapping_path = tmp_path / "mapping.json"

        config = CloakConfig(
            party_a_name="Acme",
            party_a_label="Customer",
            party_b_name="BigCo",
            party_b_label="Vendor",
        )

        result = cloak_document(input_path, requested_output, mapping_path, config)

        actual_output = Path(result.output_path)
        assert "[Customer]" in actual_output.stem
        assert "[Vendor]" in actual_output.stem
        assert "acme" not in actual_output.stem.lower() or "[customer]" in actual_output.stem.lower()
        assert actual_output.exists()

    def test_cloak_document_no_party_names_in_filename(self, tmp_path):
        """
        When the output filename does not contain party names, it is
        returned unchanged.
        """
        paragraphs = ["Agreement between Acme and BigCo."]
        input_path = make_simple_docx(tmp_path / "contract.docx", paragraphs)

        requested_output = tmp_path / "contract_cloaked.docx"
        mapping_path = tmp_path / "mapping.json"

        config = CloakConfig(
            party_a_name="Acme",
            party_a_label="Customer",
            party_b_name="BigCo",
            party_b_label="Vendor",
        )

        result = cloak_document(input_path, requested_output, mapping_path, config)

        actual_output = Path(result.output_path)
        assert actual_output.name == "contract_cloaked.docx"
        assert actual_output.exists()

    def test_cloak_document_filename_with_aliases(self, tmp_path):
        """
        Party aliases are also replaced in the output filename.
        """
        paragraphs = ['Acme Corp. ("Acme") and BigCo agree.']
        input_path = make_simple_docx(tmp_path / "Acme_BigCo_NDA.docx", paragraphs)

        requested_output = tmp_path / "Acme_BigCo_NDA_cloaked.docx"
        mapping_path = tmp_path / "mapping.json"

        config = CloakConfig(
            party_a_name="Acme Corp.",
            party_a_label="Full Vendor",
            party_b_name="BigCo",
            party_b_label="Customer",
            party_a_aliases=[PartyAlias(name="Acme", label="Vendor")],
        )

        result = cloak_document(input_path, requested_output, mapping_path, config)

        actual_output = Path(result.output_path)
        # "Acme" should be replaced (by the alias replacement) and "BigCo" by party_b.
        assert "Acme" not in actual_output.stem
        assert "BigCo" not in actual_output.stem
        assert actual_output.exists()


# ===================================================================
# _expand_content_replacements tests
# ===================================================================

class TestExpandContentReplacements:
    """Verify suffix-stripped variants are added for comment sanitization."""

    def test_adds_stripped_variant(self):
        """'Making Reign Inc.' should also generate 'Making Reign'."""
        replacements = {"Making Reign Inc.": "[Company]"}
        expanded = _expand_content_replacements(replacements)
        assert expanded["Making Reign Inc."] == "[Company]"
        assert expanded["Making Reign"] == "[Company]"

    def test_no_duplicate_if_stripped_already_present(self):
        """If the stripped form is already an explicit replacement, don't overwrite."""
        replacements = {
            "Acme Corp.": "[Vendor]",
            "Acme": "[Vendor-Short]",
        }
        expanded = _expand_content_replacements(replacements)
        assert expanded["Acme"] == "[Vendor-Short]"  # not overwritten

    def test_no_suffix_no_extra_variant(self):
        """Names without a corporate suffix produce no extra entries."""
        replacements = {"John Smith": "[Person]"}
        expanded = _expand_content_replacements(replacements)
        assert len(expanded) == 1

    def test_llc_suffix_stripped(self):
        """LLC suffix is stripped to create a variant."""
        replacements = {"Software Experts LLC": "[Vendor]"}
        expanded = _expand_content_replacements(replacements)
        assert "Software Experts" in expanded
