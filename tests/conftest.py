"""
Shared pytest fixtures for all ClientCloak tests.

Provides helper functions and fixtures for creating test .docx documents
with known content, tables, headers, footers, and comments.
"""

import json
import zipfile
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

FIXTURES_DIR = Path(__file__).parent / "fixtures"
SAMPLE_CONTRACT = FIXTURES_DIR / "sample_contract.docx"
SAMPLE_MAPPING = FIXTURES_DIR / "sample_contract_mapping.json"


# ---------------------------------------------------------------------------
# Helpers: create .docx files with known content
# ---------------------------------------------------------------------------

def make_simple_docx(path: Path, paragraphs: list[str]) -> Path:
    """Create a .docx with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


def make_table_docx(path: Path, rows: list[list[str]]) -> Path:
    """Create a .docx with a single table."""
    doc = Document()
    doc.add_paragraph("Header paragraph")
    if rows:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = cell_text
    doc.save(str(path))
    return path


def make_docx_with_comments(path: Path, body_text: str, comments: list[dict]) -> Path:
    """
    Create a .docx with body text and inject comments via ZIP manipulation.

    Each comment dict should have keys: author, initials, date, text.
    """
    # First create a basic docx
    doc = Document()
    doc.add_paragraph(body_text)
    doc.save(str(path))

    # Now inject comments.xml into the ZIP
    if comments:
        _inject_comments_xml(path, comments)

    return path


def _inject_comments_xml(path: Path, comments: list[dict]) -> None:
    """Inject a word/comments.xml entry into an existing .docx ZIP."""
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns_w)

    root = ET.Element(f"{{{ns_w}}}comments")
    for idx, c in enumerate(comments):
        comment_el = ET.SubElement(root, f"{{{ns_w}}}comment")
        comment_el.set(f"{{{ns_w}}}id", str(idx))
        comment_el.set(f"{{{ns_w}}}author", c["author"])
        comment_el.set(f"{{{ns_w}}}initials", c.get("initials", ""))
        comment_el.set(f"{{{ns_w}}}date", c.get("date", "2026-01-15T10:00:00Z"))

        p_el = ET.SubElement(comment_el, f"{{{ns_w}}}p")
        r_el = ET.SubElement(p_el, f"{{{ns_w}}}r")
        t_el = ET.SubElement(r_el, f"{{{ns_w}}}t")
        t_el.text = c["text"]

    comments_xml = ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    # Re-write the ZIP with the new comments.xml
    original = path.read_bytes()
    buf = BytesIO()
    with zipfile.ZipFile(BytesIO(original), "r") as zin, \
         zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/comments.xml":
                continue
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/comments.xml", comments_xml)

    path.write_bytes(buf.getvalue())


def make_docx_with_tracked_insertion(
    path: Path, body_text: str, inserted_text: str,
) -> Path:
    """Create a .docx with body text and a tracked insertion (w:ins) paragraph."""
    doc = Document()
    doc.add_paragraph(body_text)
    doc.save(str(path))

    # Inject a <w:ins> element via ZIP manipulation (python-docx doesn't
    # support creating tracked changes).
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    original = path.read_bytes()
    buf = BytesIO()
    with zipfile.ZipFile(BytesIO(original), "r") as zin, \
         zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            raw = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml_str = raw.decode("utf-8")
                # Insert a tracked-change paragraph before </w:body>
                ins_xml = (
                    f'<w:p><w:ins w:id="99" w:author="Test" '
                    f'w:date="2026-01-01T00:00:00Z">'
                    f'<w:r><w:t>{inserted_text}</w:t></w:r>'
                    f'</w:ins></w:p>'
                )
                xml_str = xml_str.replace("</w:body>", ins_xml + "</w:body>")
                raw = xml_str.encode("utf-8")
            zout.writestr(item, raw)
    path.write_bytes(buf.getvalue())
    return path


def make_docx_with_hidden_text(path: Path, normal_text: str, hidden_text: str) -> Path:
    """Create a .docx with a hidden run (font.hidden = True)."""
    doc = Document()
    p = doc.add_paragraph()
    run_normal = p.add_run(normal_text)
    run_hidden = p.add_run(hidden_text)
    run_hidden.font.hidden = True
    doc.save(str(path))
    return path


def make_docx_with_tiny_font(path: Path, normal_text: str, tiny_text: str) -> Path:
    """Create a .docx with a run at 1pt font size."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(normal_text)
    run_tiny = p.add_run(tiny_text)
    run_tiny.font.size = Pt(1)
    doc.save(str(path))
    return path


def make_docx_with_white_text(path: Path, normal_text: str, white_text: str) -> Path:
    """Create a .docx with a near-white colored run."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(normal_text)
    run_white = p.add_run(white_text)
    run_white.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Pytest fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def simple_docx(tmp_path):
    """A simple .docx with known text for replacement tests."""
    return make_simple_docx(
        tmp_path / "simple.docx",
        [
            "This agreement is between Acme Corporation and BigCo LLC.",
            "Acme Corporation shall provide services to BigCo LLC.",
            "Contact: jane@acme.com",
        ],
    )


@pytest.fixture
def table_docx(tmp_path):
    """A .docx with a table containing party names."""
    return make_table_docx(
        tmp_path / "table.docx",
        [
            ["Party", "Role"],
            ["Acme Corporation", "Vendor"],
            ["BigCo LLC", "Customer"],
        ],
    )


@pytest.fixture
def sample_contract():
    """Path to the pre-existing sample_contract.docx fixture."""
    assert SAMPLE_CONTRACT.exists(), f"Fixture not found: {SAMPLE_CONTRACT}"
    return SAMPLE_CONTRACT


@pytest.fixture
def sample_mapping_path():
    """Path to the pre-existing sample_contract_mapping.json fixture."""
    assert SAMPLE_MAPPING.exists(), f"Fixture not found: {SAMPLE_MAPPING}"
    return SAMPLE_MAPPING
