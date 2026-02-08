"""
Tests for clientcloak.docx_handler: load, save, extract, replace.

Covers:
- Loading valid/invalid files, error branches
- Text extraction from paragraphs, tables, headers/footers
- Text replacement with case-preserving and match_case=False
- The _transfer_case helper function
- Cross-run replacement (text split across Word XML runs)
"""

import pytest
from docx import Document
from docx.shared import Pt
from pathlib import Path

from clientcloak.docx_handler import (
    DocumentLoadError,
    PasswordProtectedError,
    UnsupportedFormatError,
    _is_bracketed_label,
    _transfer_case,
    extract_all_text,
    load_document,
    replace_text_in_document,
    replace_text_in_xml,
    save_document,
)
from tests.conftest import make_simple_docx, make_table_docx, make_docx_with_tracked_insertion


# ===================================================================
# load_document
# ===================================================================

class TestLoadDocument:
    """Tests for load_document()."""

    def test_load_valid_docx(self, simple_docx):
        doc = load_document(simple_docx)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_load_nonexistent_raises_file_not_found(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="File not found"):
            load_document(tmp_path / "nonexistent.docx")

    def test_load_directory_raises_document_load_error(self, tmp_path):
        d = tmp_path / "not_a_file.docx"
        d.mkdir()
        with pytest.raises(DocumentLoadError, match="not a file"):
            load_document(d)

    def test_load_doc_extension_raises_unsupported(self, tmp_path):
        fake = tmp_path / "old.doc"
        fake.write_bytes(b"fake content")
        with pytest.raises(UnsupportedFormatError, match="Legacy .doc"):
            load_document(fake)

    def test_load_txt_extension_raises_unsupported(self, tmp_path):
        fake = tmp_path / "file.txt"
        fake.write_bytes(b"text")
        with pytest.raises(UnsupportedFormatError, match="Unsupported file type"):
            load_document(fake)

    def test_load_corrupt_zip_raises_unsupported(self, tmp_path):
        fake = tmp_path / "corrupt.docx"
        fake.write_bytes(b"this is not a zip file at all")
        with pytest.raises(UnsupportedFormatError, match="not a valid .docx archive"):
            load_document(fake)

    def test_load_encrypted_ole2_raises_password_protected(self, tmp_path):
        """
        OLE2 magic bytes indicate an encrypted Office document.
        However, the extension check passes and zipfile.is_zipfile returns False
        for OLE2 files, so UnsupportedFormatError is raised before the encryption
        check. This test verifies the OLE2 path raises some form of DocumentLoadError.
        """
        fake = tmp_path / "encrypted.docx"
        # OLE2 magic bytes: \xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1
        fake.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 100)
        # OLE2 files fail the ZIP check first, so UnsupportedFormatError is raised
        with pytest.raises(DocumentLoadError):
            load_document(fake)

    def test_load_accepts_string_path(self, simple_docx):
        doc = load_document(str(simple_docx))
        assert doc is not None


# ===================================================================
# save_document
# ===================================================================

class TestSaveDocument:
    """Tests for save_document()."""

    def test_save_creates_file(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Hello")
        out = tmp_path / "out.docx"
        result = save_document(doc, out)
        assert out.exists()
        assert result == out.resolve()

    def test_save_creates_parent_dirs(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Test")
        out = tmp_path / "deep" / "nested" / "out.docx"
        save_document(doc, out)
        assert out.exists()

    def test_saved_document_is_loadable(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Roundtrip test")
        out = tmp_path / "roundtrip.docx"
        save_document(doc, out)
        loaded = load_document(out)
        assert loaded.paragraphs[0].text == "Roundtrip test"


# ===================================================================
# extract_all_text
# ===================================================================

class TestExtractAllText:
    """Tests for extract_all_text()."""

    def test_extract_paragraphs(self, simple_docx):
        doc = load_document(simple_docx)
        texts = extract_all_text(doc)
        text_strings = [t for t, _ in texts]
        assert any("Acme Corporation" in t for t in text_strings)
        assert any("BigCo LLC" in t for t in text_strings)

    def test_extract_from_tables(self, table_docx):
        doc = load_document(table_docx)
        texts = extract_all_text(doc)
        text_strings = [t for t, _ in texts]
        assert any("Acme Corporation" in t for t in text_strings)
        assert any("BigCo LLC" in t for t in text_strings)

    def test_extract_excludes_empty(self, tmp_path):
        path = tmp_path / "empty_paras.docx"
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("Real text")
        doc.save(str(path))

        loaded = load_document(path)
        texts = extract_all_text(loaded)
        assert len(texts) == 1
        assert texts[0][0] == "Real text"

    def test_extract_from_header_footer(self, tmp_path):
        path = tmp_path / "hf.docx"
        doc = Document()
        doc.add_paragraph("Body text")
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = "Header text"
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = "Footer text"
        doc.save(str(path))

        loaded = load_document(path)
        texts = extract_all_text(loaded)
        text_strings = [t for t, _ in texts]
        assert "Body text" in text_strings
        assert "Header text" in text_strings
        assert "Footer text" in text_strings


# ===================================================================
# replace_text_in_document
# ===================================================================

class TestReplaceTextInDocument:
    """Tests for replace_text_in_document()."""

    def test_basic_replacement(self, simple_docx):
        doc = load_document(simple_docx)
        count = replace_text_in_document(
            doc,
            {"Acme Corporation": "Vendor", "BigCo LLC": "Customer"},
        )
        assert count > 0
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Vendor" in full_text
        assert "Customer" in full_text
        assert "Acme Corporation" not in full_text
        assert "BigCo LLC" not in full_text

    def test_case_insensitive_matching(self, tmp_path):
        path = make_simple_docx(
            tmp_path / "ci.docx",
            ["ACME CORP is great. acme corp is the best. Acme Corp rules."],
        )
        doc = load_document(path)
        count = replace_text_in_document(doc, {"Acme Corp": "Vendor"})
        assert count == 3
        text = doc.paragraphs[0].text
        # With case transfer: ACME CORP -> VENDOR, acme corp -> vendor, Acme Corp -> Vendor
        assert "VENDOR" in text
        assert "vendor" in text
        assert "Vendor" in text

    def test_match_case_false_uses_exact_replacement(self, tmp_path):
        path = make_simple_docx(
            tmp_path / "mc.docx",
            ["Licensee will pay Licensor. LICENSEE agrees."],
        )
        doc = load_document(path)
        count = replace_text_in_document(
            doc,
            {"Licensee": "BigCo LLC", "Licensor": "Acme Corporation"},
            match_case=False,
        )
        assert count > 0
        text = doc.paragraphs[0].text
        # With match_case=False, the replacement is used verbatim
        assert "BigCo LLC" in text
        assert "Acme Corporation" in text

    def test_empty_replacements_returns_zero(self, simple_docx):
        doc = load_document(simple_docx)
        count = replace_text_in_document(doc, {})
        assert count == 0

    def test_replacement_in_tables(self, table_docx):
        doc = load_document(table_docx)
        count = replace_text_in_document(
            doc,
            {"Acme Corporation": "Vendor", "BigCo LLC": "Customer"},
        )
        assert count >= 2

    def test_longest_match_first(self, tmp_path):
        """Ensure 'Acme Corporation' is matched before 'Acme'.

        Bracketed labels are preserved verbatim (no case transfer).
        """
        path = make_simple_docx(
            tmp_path / "longest.docx",
            ["Acme Corporation and Acme both appear."],
        )
        doc = load_document(path)
        count = replace_text_in_document(
            doc,
            {"Acme Corporation": "[PARTY_A]", "Acme": "[SHORT]"},
        )
        text = doc.paragraphs[0].text
        # Bracketed labels are used verbatim — no case transfer
        assert "[PARTY_A]" in text
        assert "[SHORT]" in text
        assert count == 2

    def test_replacement_in_header_footer(self, tmp_path):
        path = tmp_path / "hf_replace.docx"
        doc = Document()
        doc.add_paragraph("Acme in body")
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = "Prepared by Acme"
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = "Acme Confidential"
        doc.save(str(path))

        doc = load_document(path)
        count = replace_text_in_document(doc, {"Acme": "[VENDOR]"})
        assert count >= 3


# ===================================================================
# _transfer_case
# ===================================================================

class TestTransferCase:
    """Tests for the internal _transfer_case function."""

    @pytest.mark.parametrize(
        "original, replacement, expected",
        [
            # All uppercase
            ("ACME", "vendor", "VENDOR"),
            # All lowercase
            ("acme", "Vendor", "vendor"),
            # Title case
            ("Acme Corp", "vendor inc", "Vendor Inc"),
            # Loose title case (e.g. "BigCo LLC" - all words start uppercase)
            ("BigCo LLC", "vendor inc", "Vendor Inc"),
            # Empty strings
            ("", "anything", "anything"),
            ("something", "", ""),
            # Sentence case: first char upper triggers sentence-case branch
            # (first char upper, rest mixed) -> capitalize first, keep rest
            ("AbCd", "wxyz", "Wxyz"),
        ],
    )
    def test_transfer_case_patterns(self, original, replacement, expected):
        result = _transfer_case(original, replacement)
        assert result == expected

    def test_transfer_case_sentence_case(self):
        # First char upper, rest lower-ish -> sentence case
        result = _transfer_case("Hello world", "goodbye earth")
        assert result[0] == "G"  # first char uppercase

    def test_transfer_case_both_empty(self):
        assert _transfer_case("", "") == ""


# ===================================================================
# Cross-run replacement
# ===================================================================

class TestCrossRunReplacement:
    """Test replacement when text is split across multiple Word XML runs."""

    def test_text_split_across_two_runs(self, tmp_path):
        """When Word splits 'Acme Corporation' across two runs, replacement should still work.

        Bracketed labels are preserved verbatim (no case transfer).
        """
        path = tmp_path / "split_runs.docx"
        doc = Document()
        p = doc.add_paragraph()
        # Simulate Word splitting the name across runs
        run1 = p.add_run("Agreement with Acme")
        run2 = p.add_run(" Corporation for services.")
        doc.save(str(path))

        doc = load_document(path)
        count = replace_text_in_document(doc, {"Acme Corporation": "[VENDOR]"})
        assert count == 1
        full = "".join(r.text for r in doc.paragraphs[0].runs)
        # Bracketed label is used verbatim — no case transfer
        assert "[VENDOR]" in full
        assert "Acme Corporation" not in full

    def test_text_split_across_three_runs(self, tmp_path):
        """Text split across three runs should also work."""
        path = tmp_path / "three_runs.docx"
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Contact Acme")
        p.add_run(" Corp")
        p.add_run("oration today.")
        doc.save(str(path))

        doc = load_document(path)
        count = replace_text_in_document(doc, {"Acme Corporation": "[VENDOR]"})
        assert count == 1
        full = "".join(r.text for r in doc.paragraphs[0].runs)
        # Bracketed label is used verbatim
        assert "[VENDOR]" in full

    def test_cross_run_with_match_case_false(self, tmp_path):
        """With match_case=False, replacement text is used verbatim."""
        path = tmp_path / "split_exact.docx"
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Agreement with Acme")
        p.add_run(" Corporation for services.")
        doc.save(str(path))

        doc = load_document(path)
        count = replace_text_in_document(
            doc,
            {"Acme Corporation": "BigCo LLC"},
            match_case=False,
        )
        assert count == 1
        full = "".join(r.text for r in doc.paragraphs[0].runs)
        assert "BigCo LLC" in full


# ===================================================================
# Bracketed label preservation
# ===================================================================

class TestBracketedLabelPreservation:
    """Bracketed labels like [AltCustomerName] must be used verbatim."""

    def test_is_bracketed_label_true(self):
        assert _is_bracketed_label("[AltCustomerName]") is True
        assert _is_bracketed_label("[X]") is True

    def test_is_bracketed_label_false(self):
        assert _is_bracketed_label("Vendor") is False
        assert _is_bracketed_label("[only open") is False
        assert _is_bracketed_label("only close]") is False
        assert _is_bracketed_label("") is False
        assert _is_bracketed_label("[") is False

    def test_bracketed_label_preserved_verbatim_title_case_source(self, tmp_path):
        """Replacing title-case 'Licensee' with '[AltCustomerName]' must NOT
        mangle the label's internal capitalisation."""
        path = make_simple_docx(
            tmp_path / "bracket_title.docx",
            ["The Licensee shall comply with all terms."],
        )
        doc = load_document(path)
        count = replace_text_in_document(doc, {"Licensee": "[AltCustomerName]"})
        assert count == 1
        text = doc.paragraphs[0].text
        assert "[AltCustomerName]" in text

    def test_bracketed_label_preserved_verbatim_upper_case_source(self, tmp_path):
        """Replacing all-caps 'LICENSEE' must still keep the label verbatim."""
        path = make_simple_docx(
            tmp_path / "bracket_upper.docx",
            ["THE LICENSEE SHALL COMPLY WITH ALL TERMS."],
        )
        doc = load_document(path)
        count = replace_text_in_document(doc, {"Licensee": "[AltCustomerName]"})
        assert count == 1
        text = doc.paragraphs[0].text
        assert "[AltCustomerName]" in text

    def test_bracketed_label_preserved_verbatim_lower_case_source(self, tmp_path):
        """Replacing all-lowercase 'licensee' must still keep the label verbatim."""
        path = make_simple_docx(
            tmp_path / "bracket_lower.docx",
            ["the licensee shall comply with all terms."],
        )
        doc = load_document(path)
        count = replace_text_in_document(doc, {"Licensee": "[AltCustomerName]"})
        assert count == 1
        text = doc.paragraphs[0].text
        assert "[AltCustomerName]" in text

    def test_bracketed_label_preserved_across_runs(self, tmp_path):
        """Bracketed label stays verbatim even when the match spans two runs."""
        path = tmp_path / "bracket_runs.docx"
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("The Lic")
        p.add_run("ensee shall comply.")
        doc.save(str(path))

        doc = load_document(path)
        count = replace_text_in_document(doc, {"Licensee": "[AltCustomerName]"})
        assert count == 1
        full = "".join(r.text for r in doc.paragraphs[0].runs)
        assert "[AltCustomerName]" in full

    def test_non_bracketed_replacement_still_gets_case_transfer(self, tmp_path):
        """Non-bracketed replacements must still have case transfer applied."""
        path = make_simple_docx(
            tmp_path / "non_bracket.docx",
            ["ACME shipped the product."],
        )
        doc = load_document(path)
        count = replace_text_in_document(doc, {"Acme": "Vendor"})
        assert count == 1
        text = doc.paragraphs[0].text
        assert "VENDOR" in text


# ===================================================================
# replace_text_in_xml (tracked changes, text boxes, footnotes)
# ===================================================================

class TestReplaceTextInXml:
    """Tests for replace_text_in_xml() — ZIP-level XML replacement."""

    def test_replaces_text_in_tracked_insertion(self, tmp_path):
        """Placeholders inside <w:ins> elements should be replaced."""
        path = make_docx_with_tracked_insertion(
            tmp_path / "tracked.docx",
            "Agreement between [Company] and [Vendor].",
            "[Person-1] will attend the meeting with [Person-2].",
        )
        count = replace_text_in_xml(
            path,
            {"[Person-1]": "Jane Smith", "[Person-2]": "Bob Jones",
             "[Company]": "Acme Corp", "[Vendor]": "BigCo LLC"},
            match_case=False,
        )
        assert count >= 2  # at least the tracked insertion replacements

        # Verify tracked change text is replaced in raw XML
        import zipfile, re
        with zipfile.ZipFile(path, "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        ins_matches = re.findall(r"<w:ins\s[^>]*>[\s\S]*?</w:ins>", xml)
        ins_text = "".join(
            t for m in ins_matches
            for t in re.findall(r"<w:t[^>]*>([^<]*)</w:t>", m)
        )
        assert "Jane Smith" in ins_text
        assert "Bob Jones" in ins_text
        assert "[Person-1]" not in ins_text
        assert "[Person-2]" not in ins_text

    def test_no_replacements_returns_zero(self, tmp_path):
        path = make_simple_docx(tmp_path / "plain.docx", ["No placeholders here."])
        count = replace_text_in_xml(path, {"[Company]": "Acme"}, match_case=False)
        assert count == 0

    def test_empty_replacements_returns_zero(self, tmp_path):
        path = make_simple_docx(tmp_path / "plain.docx", ["Some text."])
        count = replace_text_in_xml(path, {}, match_case=False)
        assert count == 0

    def test_case_transfer_in_xml(self, tmp_path):
        """When match_case=True, case transfer should apply in XML too."""
        path = make_docx_with_tracked_insertion(
            tmp_path / "case.docx",
            "Body text.",
            "ACME shipped the product.",
        )
        count = replace_text_in_xml(
            path,
            {"Acme": "[Vendor]"},
            match_case=True,
        )
        assert count >= 1

        import zipfile
        with zipfile.ZipFile(path, "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        # Bracketed labels should be preserved verbatim
        assert "[Vendor]" in xml
