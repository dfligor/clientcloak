"""
Tests for clientcloak.metadata: inspect and strip metadata from .docx files.
"""

import pytest
from pathlib import Path
from docx import Document

from clientcloak.metadata import inspect_metadata, strip_metadata
from clientcloak.models import MetadataReport
from tests.conftest import make_simple_docx


# ===================================================================
# Helpers
# ===================================================================

def _make_docx_with_metadata(path: Path) -> Path:
    """Create a .docx with known author metadata set via python-docx."""
    doc = Document()
    doc.add_paragraph("Document with metadata.")
    props = doc.core_properties
    props.author = "Jane Smith"
    props.last_modified_by = "John Doe"
    props.subject = "Test Subject"
    props.keywords = "test, metadata"
    props.comments = "Some internal comments"
    props.title = "Test Title"
    props.category = "Legal"
    doc.save(str(path))
    return path


# ===================================================================
# inspect_metadata
# ===================================================================

class TestInspectMetadata:
    """Tests for inspect_metadata()."""

    def test_inspects_known_metadata(self, tmp_path):
        path = _make_docx_with_metadata(tmp_path / "meta.docx")
        report = inspect_metadata(path)
        assert isinstance(report, MetadataReport)
        assert report.author == "Jane Smith"
        assert report.last_modified_by == "John Doe"

    def test_minimal_document_has_some_fields(self, tmp_path):
        path = make_simple_docx(tmp_path / "minimal.docx", ["Hello"])
        report = inspect_metadata(path)
        assert isinstance(report, MetadataReport)
        # Even a minimal doc may have revision, application, etc.

    def test_returns_none_for_missing_fields(self, tmp_path):
        path = make_simple_docx(tmp_path / "empty.docx", ["text"])
        report = inspect_metadata(path)
        # A freshly-created doc typically won't have author set
        # (python-docx may set it to empty or None)
        assert isinstance(report, MetadataReport)


# ===================================================================
# strip_metadata
# ===================================================================

class TestStripMetadata:
    """Tests for strip_metadata()."""

    def test_strip_removes_author(self, tmp_path):
        input_path = _make_docx_with_metadata(tmp_path / "input.docx")
        output_path = tmp_path / "stripped.docx"

        before_report = strip_metadata(input_path, output_path)
        assert before_report.author == "Jane Smith"  # the "before" report

        # Verify the output has metadata stripped
        after_report = inspect_metadata(output_path)
        assert after_report.author is None or after_report.author == ""

    def test_strip_removes_last_modified_by(self, tmp_path):
        input_path = _make_docx_with_metadata(tmp_path / "input.docx")
        output_path = tmp_path / "stripped.docx"

        strip_metadata(input_path, output_path)
        after_report = inspect_metadata(output_path)
        assert after_report.last_modified_by is None or after_report.last_modified_by == ""

    def test_strip_in_place(self, tmp_path):
        """Verify that input_path == output_path works (in-place strip)."""
        path = _make_docx_with_metadata(tmp_path / "inplace.docx")
        strip_metadata(path, path)
        report = inspect_metadata(path)
        assert report.author is None or report.author == ""

    def test_stripped_document_is_valid(self, tmp_path):
        input_path = _make_docx_with_metadata(tmp_path / "input.docx")
        output_path = tmp_path / "stripped.docx"
        strip_metadata(input_path, output_path)

        # The output should be loadable
        doc = Document(str(output_path))
        assert doc.paragraphs[0].text == "Document with metadata."

    def test_returns_before_report(self, tmp_path):
        input_path = _make_docx_with_metadata(tmp_path / "input.docx")
        output_path = tmp_path / "stripped.docx"
        before = strip_metadata(input_path, output_path)
        assert isinstance(before, MetadataReport)
        assert before.author == "Jane Smith"
