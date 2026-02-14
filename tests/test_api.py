"""
Integration tests for the FastAPI web API routes.

Uses httpx.AsyncClient (via pytest-asyncio is not required — we use the
synchronous TestClient from Starlette) to exercise the upload, cloak,
download, uncloak, and download-uncloaked endpoints.
"""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from clientcloak.ui.app import create_app


def _parse_upload_ndjson(resp):
    """Parse an NDJSON streaming upload response, returning the final data dict."""
    lines = resp.text.strip().split("\n")
    for line in reversed(lines):
        event = json.loads(line)
        if event.get("stage") == "complete":
            return event["data"]
    raise ValueError("No 'complete' event in upload response")


@pytest.fixture()
def client(tmp_path, monkeypatch):
    """Create a TestClient with a temporary session directory."""
    # Point sessions to a temp directory so tests don't pollute real data.
    monkeypatch.setenv("CLIENTCLOAK_SESSIONS_DIR", str(tmp_path / "sessions"))
    app = create_app()
    return TestClient(app)


@pytest.fixture()
def sample_docx(tmp_path) -> Path:
    """Create a minimal .docx file for upload tests."""
    doc = Document()
    doc.add_paragraph("This agreement is between Acme Corporation and BigCo LLC.")
    doc.add_paragraph("Acme Corporation shall provide services to BigCo LLC.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Upload
# ---------------------------------------------------------------------------


def test_upload_success(client, sample_docx):
    with open(sample_docx, "rb") as f:
        resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    assert resp.status_code == 200
    data = _parse_upload_ndjson(resp)
    assert "session_id" in data
    assert data["filename"] == "test.docx"


def test_upload_rejects_non_docx(client, tmp_path):
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("not a docx")
    with open(txt_file, "rb") as f:
        resp = client.post("/api/upload", files={"file": ("test.txt", f)})
    assert resp.status_code == 400
    assert "Only .docx" in resp.json()["detail"]


def test_upload_rejects_missing_filename(client):
    resp = client.post("/api/upload", files={"file": ("", b"")})
    # FastAPI may reject at the validation layer (422) or the handler (400).
    assert resp.status_code in (400, 422)


# ---------------------------------------------------------------------------
# Cloak
# ---------------------------------------------------------------------------


def test_cloak_success(client, sample_docx):
    # Upload first
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]

    # Cloak
    resp = client.post("/api/cloak", data={
        "session_id": session_id,
        "party_a": "Acme Corporation",
        "party_b": "BigCo LLC",
        "party_a_label": "Customer",
        "party_b_label": "Vendor",
        "comment_mode": "strip",
        "strip_metadata": True,
    })
    assert resp.status_code == 200
    data = resp.json()
    assert data["replacements_applied"] > 0
    assert "download_url" in data
    assert "mapping_url" in data


def test_cloak_invalid_session(client):
    resp = client.post("/api/cloak", data={
        "session_id": "nonexistent",
        "party_a": "A",
        "party_b": "B",
    })
    assert resp.status_code == 404


def test_cloak_invalid_comment_mode(client, sample_docx):
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]

    resp = client.post("/api/cloak", data={
        "session_id": session_id,
        "party_a": "A",
        "party_b": "B",
        "comment_mode": "invalid",
    })
    assert resp.status_code == 400
    assert "Invalid comment_mode" in resp.json()["detail"]


# ---------------------------------------------------------------------------
# Download
# ---------------------------------------------------------------------------


def test_download_cloaked(client, sample_docx):
    # Upload + cloak
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]
    client.post("/api/cloak", data={
        "session_id": session_id,
        "party_a": "Acme Corporation",
        "party_b": "BigCo LLC",
    })

    # Download cloaked
    resp = client.get(f"/api/download/{session_id}/cloaked")
    assert resp.status_code == 200
    assert "application/vnd.openxmlformats" in resp.headers["content-type"]


def test_download_mapping(client, sample_docx):
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]
    client.post("/api/cloak", data={
        "session_id": session_id,
        "party_a": "Acme Corporation",
        "party_b": "BigCo LLC",
    })

    resp = client.get(f"/api/download/{session_id}/mapping")
    assert resp.status_code == 200
    mapping = json.loads(resp.content)
    assert "mappings" in mapping


def test_download_invalid_file_type(client, sample_docx):
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]

    resp = client.get(f"/api/download/{session_id}/invalid")
    assert resp.status_code == 400


def test_download_before_cloak(client, sample_docx):
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]

    resp = client.get(f"/api/download/{session_id}/cloaked")
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Uncloak
# ---------------------------------------------------------------------------


def test_uncloak_roundtrip(client, sample_docx):
    # Upload + cloak
    with open(sample_docx, "rb") as f:
        upload_resp = client.post("/api/upload", files={"file": ("test.docx", f)})
    session_id = _parse_upload_ndjson(upload_resp)["session_id"]
    cloak_resp = client.post("/api/cloak", data={
        "session_id": session_id,
        "party_a": "Acme Corporation",
        "party_b": "BigCo LLC",
        "party_a_label": "Customer",
        "party_b_label": "Vendor",
    })

    # Download cloaked doc and mapping
    cloaked_bytes = client.get(f"/api/download/{session_id}/cloaked").content
    mapping_bytes = client.get(f"/api/download/{session_id}/mapping").content

    # Uncloak
    resp = client.post("/api/uncloak", files={
        "redlined_file": ("cloaked.docx", BytesIO(cloaked_bytes)),
        "mapping_file": ("mapping.json", BytesIO(mapping_bytes)),
    })
    assert resp.status_code == 200
    data = resp.json()
    assert data["replacements_restored"] > 0
    assert "download_url" in data


def test_uncloak_rejects_non_docx(client, tmp_path):
    mapping = tmp_path / "map.json"
    mapping.write_text('{"mappings": {}}')
    txt = tmp_path / "test.txt"
    txt.write_text("not docx")

    with open(txt, "rb") as rf, open(mapping, "rb") as mf:
        resp = client.post("/api/uncloak", files={
            "redlined_file": ("test.txt", rf),
            "mapping_file": ("map.json", mf),
        })
    assert resp.status_code == 400


def test_uncloak_rejects_non_json_mapping(client, sample_docx):
    with open(sample_docx, "rb") as rf:
        resp = client.post("/api/uncloak", files={
            "redlined_file": ("test.docx", rf),
            "mapping_file": ("map.txt", BytesIO(b"{}")),
        })
    assert resp.status_code == 400


# ---------------------------------------------------------------------------
# Index
# ---------------------------------------------------------------------------


def test_index_returns_html(client):
    resp = client.get("/")
    assert resp.status_code == 200
    assert "text/html" in resp.headers["content-type"]
